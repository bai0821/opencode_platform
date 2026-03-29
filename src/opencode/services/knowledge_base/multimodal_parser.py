"""
多模態文件解析器

支援：
- PDF（含圖片 OCR、表格）
- Word (.docx)
- Excel (.xlsx, .csv)
- Markdown (.md)
- 純文字 (.txt)
- JSON (.json)
- 程式碼 (.py, .js, .ts, .java, .cpp 等)

圖片處理：使用 Docling OCR 提取圖片中的文字
"""

import os
import io
import json
import logging
from typing import List, Dict, Any, Optional
from pathlib import Path

logger = logging.getLogger(__name__)

# 支援的文件格式
SUPPORTED_FORMATS = {
    # 文檔類
    '.pdf': 'pdf',
    '.docx': 'word',
    '.doc': 'word',
    '.xlsx': 'excel',
    '.xls': 'excel',
    '.csv': 'csv',
    '.tsv': 'csv',
    
    # 文字類
    '.txt': 'text',
    '.md': 'markdown',
    '.markdown': 'markdown',
    '.rst': 'text',
    
    # 數據類
    '.json': 'json',
    '.jsonl': 'json',
    '.yaml': 'yaml',
    '.yml': 'yaml',
    '.xml': 'xml',
    
    # 程式碼類
    '.py': 'code',
    '.js': 'code',
    '.ts': 'code',
    '.jsx': 'code',
    '.tsx': 'code',
    '.java': 'code',
    '.cpp': 'code',
    '.c': 'code',
    '.h': 'code',
    '.hpp': 'code',
    '.cs': 'code',
    '.go': 'code',
    '.rs': 'code',
    '.rb': 'code',
    '.php': 'code',
    '.swift': 'code',
    '.kt': 'code',
    '.scala': 'code',
    '.r': 'code',
    '.sql': 'code',
    '.sh': 'code',
    '.bash': 'code',
    '.ps1': 'code',
    '.html': 'code',
    '.css': 'code',
    '.scss': 'code',
    '.less': 'code',
}


class MultimodalParser:
    """
    多模態文件解析器
    
    特色：
    1. 自動識別文件類型
    2. PyMuPDF + Tesseract OCR 提取圖片文字
    3. 表格結構化提取
    4. 支援 15+ 種文件格式
    """
    
    def __init__(
        self,
        chunk_size: int = 512,
        chunk_overlap: int = 50,
        enable_ocr: bool = True
    ):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.enable_ocr = enable_ocr
        
        # 檢查 OCR 工具
        self._ocr_available = False
        self._init_ocr()
    
    def _init_ocr(self):
        """初始化 OCR 工具"""
        # 方法1：嘗試使用 pytesseract
        try:
            import pytesseract
            from PIL import Image
            # 測試 tesseract 是否可用
            pytesseract.get_tesseract_version()
            self._ocr_available = True
            self._ocr_method = "pytesseract"
            logger.info("✅ [MultimodalParser] Tesseract OCR 初始化成功")
        except Exception as e:
            logger.warning(f"⚠️ [MultimodalParser] Tesseract OCR 不可用: {e}")
            
            # 方法2：嘗試使用 easyocr
            try:
                import easyocr
                self._ocr_reader = easyocr.Reader(['ch_tra', 'en'], gpu=False)
                self._ocr_available = True
                self._ocr_method = "easyocr"
                logger.info("✅ [MultimodalParser] EasyOCR 初始化成功")
            except Exception as e2:
                logger.warning(f"⚠️ [MultimodalParser] EasyOCR 也不可用: {e2}")
                logger.info("💡 提示：安裝 OCR 工具來提取圖片文字")
                logger.info("   pip install pytesseract pillow")
                logger.info("   或 pip install easyocr")
    
    def _ocr_image(self, image_bytes: bytes) -> str:
        """對圖片進行 OCR"""
        if not self._ocr_available:
            return ""
        
        try:
            if self._ocr_method == "pytesseract":
                import pytesseract
                from PIL import Image
                import io
                
                image = Image.open(io.BytesIO(image_bytes))
                # 使用中文+英文識別
                text = pytesseract.image_to_string(image, lang='chi_tra+eng')
                return text.strip()
                
            elif self._ocr_method == "easyocr":
                import numpy as np
                from PIL import Image
                import io
                
                image = Image.open(io.BytesIO(image_bytes))
                image_np = np.array(image)
                results = self._ocr_reader.readtext(image_np)
                text = " ".join([r[1] for r in results])
                return text.strip()
                
        except Exception as e:
            logger.warning(f"⚠️ [OCR] 識別失敗: {e}")
            return ""
        
        return ""
    
    def parse(self, file_path: str) -> List[Dict[str, Any]]:
        """
        解析文件（自動識別類型）
        
        Args:
            file_path: 文件路徑
            
        Returns:
            解析後的 chunks 列表
        """
        if not os.path.exists(file_path):
            logger.error(f"❌ [Parser] 文件不存在: {file_path}")
            return []
        
        ext = Path(file_path).suffix.lower()
        file_type = SUPPORTED_FORMATS.get(ext)
        
        if not file_type:
            logger.warning(f"⚠️ [Parser] 不支援的文件格式: {ext}")
            file_type = 'text'
        
        logger.info(f"📄 [Parser] 解析文件: {file_path} (類型: {file_type})")
        
        # 根據類型調用對應的解析器
        parsers = {
            'pdf': self._parse_pdf,
            'word': self._parse_word,
            'excel': self._parse_excel,
            'csv': self._parse_csv,
            'text': self._parse_text,
            'markdown': self._parse_markdown,
            'json': self._parse_json,
            'yaml': self._parse_yaml,
            'xml': self._parse_xml,
            'code': self._parse_code,
        }
        
        parser_func = parsers.get(file_type, self._parse_text)
        return parser_func(file_path)
    
    def _parse_pdf(self, file_path: str) -> List[Dict[str, Any]]:
        """
        解析 PDF - 使用 PyMuPDF，支援圖片 OCR
        """
        try:
            import fitz  # PyMuPDF
        except ImportError:
            logger.error("❌ [Parser] PyMuPDF 未安裝: pip install pymupdf")
            return []
        
        parsed_data = []
        file_name = os.path.basename(file_path)
        
        try:
            doc = fitz.open(file_path)
            total_pages = len(doc)
            logger.info(f"📖 [Parser] PDF 共 {total_pages} 頁")
            
            for page_num in range(total_pages):
                page = doc[page_num]
                # 頁碼從 1 開始
                page_label = str(page_num + 1)
                
                logger.info(f"📄 [Parser] 解析第 {page_label}/{total_pages} 頁...")
                
                # 1. 提取文字內容
                text = page.get_text("text").strip()
                if text and len(text) > 30:
                    chunks = self._split_text(text)
                    for chunk_idx, chunk in enumerate(chunks):
                        if len(chunk) > 30:
                            parsed_data.append({
                                "text": chunk,
                                "metadata": {
                                    "file_name": file_name,
                                    "page_label": page_label,
                                    "chunk_index": chunk_idx,
                                    "content_type": "text",
                                    "source": "pymupdf"
                                }
                            })
                    logger.info(f"   文字: {len(chunks)} 個區塊")
                
                # 2. 提取圖片並進行 OCR
                if self.enable_ocr:
                    images = page.get_images(full=True)
                    image_count = 0
                    
                    for img_idx, img in enumerate(images):
                        try:
                            xref = img[0]
                            base_image = doc.extract_image(xref)
                            image_bytes = base_image["image"]
                            
                            # 跳過太小的圖片（可能是圖標）
                            if len(image_bytes) < 3000:  # 小於 3KB
                                continue
                            
                            # 進行 OCR
                            ocr_text = self._ocr_image(image_bytes)
                            
                            if ocr_text and len(ocr_text) > 10:
                                parsed_data.append({
                                    "text": f"[圖片內容 - 第 {page_label} 頁 圖 {img_idx + 1}]\n{ocr_text}",
                                    "metadata": {
                                        "file_name": file_name,
                                        "page_label": page_label,
                                        "content_type": "image_ocr",
                                        "image_index": img_idx,
                                        "source": "ocr"
                                    }
                                })
                                image_count += 1
                                
                        except Exception as e:
                            logger.debug(f"圖片 {img_idx} 處理失敗: {e}")
                    
                    if image_count > 0:
                        logger.info(f"   圖片 OCR: {image_count} 張")
                
                # 3. 提取表格
                try:
                    tables = page.find_tables()
                    if tables and len(tables.tables) > 0:
                        for table_idx, table in enumerate(tables.tables):
                            table_data = table.extract()
                            if table_data:
                                table_text = self._format_table(table_data)
                                if table_text and len(table_text) > 30:
                                    parsed_data.append({
                                        "text": f"[表格 - 第 {page_label} 頁 表 {table_idx + 1}]\n{table_text}",
                                        "metadata": {
                                            "file_name": file_name,
                                            "page_label": page_label,
                                            "content_type": "table",
                                            "table_index": table_idx,
                                            "source": "pymupdf_table"
                                        }
                                    })
                        if len(tables.tables) > 0:
                            logger.info(f"   表格: {len(tables.tables)} 個")
                except Exception as e:
                    logger.debug(f"表格提取失敗: {e}")
            
            doc.close()
            
            # 統計
            content_types = {}
            for item in parsed_data:
                ct = item.get("metadata", {}).get("content_type", "text")
                content_types[ct] = content_types.get(ct, 0) + 1
            
            summary = ", ".join([f"{k}: {v}" for k, v in content_types.items()])
            logger.info(f"✅ [Parser] PDF 解析完成: {len(parsed_data)} 個區塊 ({summary})")
            
        except Exception as e:
            logger.error(f"❌ [Parser] PDF 解析失敗: {e}")
            import traceback
            logger.error(traceback.format_exc())
        
        return parsed_data
    
    def _format_table(self, table_data: List[List]) -> str:
        """將表格數據格式化為文字"""
        if not table_data:
            return ""
        
        lines = []
        for row in table_data:
            row_text = " | ".join([str(cell) if cell else "" for cell in row])
            if row_text.strip():
                lines.append(row_text)
        
        return "\n".join(lines)
    
    def _parse_word(self, file_path: str) -> List[Dict[str, Any]]:
        """解析 Word 文檔"""
        try:
            from docx import Document
        except ImportError:
            logger.error("❌ [Parser] python-docx 未安裝: pip install python-docx")
            return []
        
        parsed_data = []
        file_name = os.path.basename(file_path)
        
        try:
            doc = Document(file_path)
            full_text = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text.strip())
            
            # 處理表格
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        table_text.append(row_text)
                if table_text:
                    full_text.append("[表格]\n" + "\n".join(table_text))
            
            combined_text = "\n\n".join(full_text)
            chunks = self._split_text(combined_text)
            
            for chunk_idx, chunk in enumerate(chunks):
                if len(chunk) > 30:
                    parsed_data.append({
                        "text": chunk,
                        "metadata": {
                            "file_name": file_name,
                            "page_label": "1",
                            "chunk_index": chunk_idx,
                            "content_type": "text",
                            "source": "python-docx"
                        }
                    })
            
            logger.info(f"✅ [Parser] Word 解析完成，共 {len(parsed_data)} 個 chunks")
            
        except Exception as e:
            logger.error(f"❌ [Parser] Word 解析失敗: {e}")
        
        return parsed_data
    
    def _parse_excel(self, file_path: str) -> List[Dict[str, Any]]:
        """解析 Excel 文件"""
        try:
            import pandas as pd
        except ImportError:
            logger.error("❌ [Parser] pandas 未安裝: pip install pandas openpyxl")
            return []
        
        parsed_data = []
        file_name = os.path.basename(file_path)
        
        try:
            xlsx = pd.ExcelFile(file_path)
            
            for sheet_name in xlsx.sheet_names:
                df = pd.read_excel(xlsx, sheet_name=sheet_name)
                
                sheet_text = f"[工作表: {sheet_name}]\n"
                sheet_text += f"欄位: {', '.join(df.columns.astype(str))}\n"
                sheet_text += f"資料筆數: {len(df)}\n\n"
                
                if len(df) > 0:
                    sample = df.head(10).to_string()
                    sheet_text += f"數據範例:\n{sample}"
                
                try:
                    desc = df.describe().to_string()
                    sheet_text += f"\n\n統計摘要:\n{desc}"
                except Exception as e:
                    logger.warning(f"⚠️ [Parser] 生成統計摘要失敗: {e}")
                
                chunks = self._split_text(sheet_text)
                for chunk_idx, chunk in enumerate(chunks):
                    if len(chunk) > 30:
                        parsed_data.append({
                            "text": chunk,
                            "metadata": {
                                "file_name": file_name,
                                "page_label": sheet_name,
                                "chunk_index": chunk_idx,
                                "content_type": "spreadsheet",
                                "source": "pandas"
                            }
                        })
            
            logger.info(f"✅ [Parser] Excel 解析完成，共 {len(parsed_data)} 個 chunks")
            
        except Exception as e:
            logger.error(f"❌ [Parser] Excel 解析失敗: {e}")
        
        return parsed_data
    
    def _parse_csv(self, file_path: str) -> List[Dict[str, Any]]:
        """解析 CSV 文件"""
        try:
            import pandas as pd
        except ImportError:
            logger.error("❌ [Parser] pandas 未安裝")
            return []
        
        parsed_data = []
        file_name = os.path.basename(file_path)
        
        try:
            for encoding in ['utf-8', 'utf-8-sig', 'gbk', 'big5', 'latin1']:
                try:
                    df = pd.read_csv(file_path, encoding=encoding)
                    break
                except Exception as e:
                    logger.warning(f"⚠️ [Parser] 以 {encoding} 編碼讀取 CSV 失敗: {e}")
                    continue
            else:
                logger.error(f"❌ [Parser] 無法讀取 CSV 文件")
                return []
            
            csv_text = f"[CSV 文件: {file_name}]\n"
            csv_text += f"欄位: {', '.join(df.columns.astype(str))}\n"
            csv_text += f"資料筆數: {len(df)}\n\n"
            
            if len(df) > 0:
                sample = df.head(20).to_string()
                csv_text += f"數據範例:\n{sample}"
            
            chunks = self._split_text(csv_text)
            for chunk_idx, chunk in enumerate(chunks):
                if len(chunk) > 30:
                    parsed_data.append({
                        "text": chunk,
                        "metadata": {
                            "file_name": file_name,
                            "page_label": "1",
                            "chunk_index": chunk_idx,
                            "content_type": "csv",
                            "source": "pandas"
                        }
                    })
            
            logger.info(f"✅ [Parser] CSV 解析完成，共 {len(parsed_data)} 個 chunks")
            
        except Exception as e:
            logger.error(f"❌ [Parser] CSV 解析失敗: {e}")
        
        return parsed_data
    
    def _parse_text(self, file_path: str) -> List[Dict[str, Any]]:
        """解析純文字文件"""
        parsed_data = []
        file_name = os.path.basename(file_path)
        
        try:
            content = None
            for encoding in ['utf-8', 'utf-8-sig', 'gbk', 'big5', 'latin1']:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    break
                except Exception as e:
                    logger.warning(f"⚠️ [Parser] 以 {encoding} 編碼讀取文字文件失敗: {e}")
                    continue
            
            if not content:
                logger.error(f"❌ [Parser] 無法讀取文件")
                return []
            
            chunks = self._split_text(content)
            for chunk_idx, chunk in enumerate(chunks):
                if len(chunk) > 30:
                    parsed_data.append({
                        "text": chunk,
                        "metadata": {
                            "file_name": file_name,
                            "page_label": "1",
                            "chunk_index": chunk_idx,
                            "content_type": "text",
                            "source": "text"
                        }
                    })
            
            logger.info(f"✅ [Parser] 文字文件解析完成，共 {len(parsed_data)} 個 chunks")
            
        except Exception as e:
            logger.error(f"❌ [Parser] 文字文件解析失敗: {e}")
        
        return parsed_data
    
    def _parse_markdown(self, file_path: str) -> List[Dict[str, Any]]:
        """解析 Markdown 文件"""
        return self._parse_text(file_path)
    
    def _parse_json(self, file_path: str) -> List[Dict[str, Any]]:
        """解析 JSON 文件"""
        parsed_data = []
        file_name = os.path.basename(file_path)
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            json_text = json.dumps(data, ensure_ascii=False, indent=2)
            
            chunks = self._split_text(json_text)
            for chunk_idx, chunk in enumerate(chunks):
                if len(chunk) > 30:
                    parsed_data.append({
                        "text": chunk,
                        "metadata": {
                            "file_name": file_name,
                            "page_label": "1",
                            "chunk_index": chunk_idx,
                            "content_type": "json",
                            "source": "json"
                        }
                    })
            
            logger.info(f"✅ [Parser] JSON 解析完成，共 {len(parsed_data)} 個 chunks")
            
        except Exception as e:
            logger.error(f"❌ [Parser] JSON 解析失敗: {e}")
        
        return parsed_data
    
    def _parse_yaml(self, file_path: str) -> List[Dict[str, Any]]:
        """解析 YAML 文件"""
        try:
            import yaml
        except ImportError:
            logger.warning("⚠️ [Parser] PyYAML 未安裝，作為文字處理")
            return self._parse_text(file_path)
        
        parsed_data = []
        file_name = os.path.basename(file_path)
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = yaml.safe_load(f)
            
            yaml_text = yaml.dump(data, allow_unicode=True, default_flow_style=False)
            
            chunks = self._split_text(yaml_text)
            for chunk_idx, chunk in enumerate(chunks):
                if len(chunk) > 30:
                    parsed_data.append({
                        "text": chunk,
                        "metadata": {
                            "file_name": file_name,
                            "page_label": "1",
                            "chunk_index": chunk_idx,
                            "content_type": "yaml",
                            "source": "yaml"
                        }
                    })
            
            logger.info(f"✅ [Parser] YAML 解析完成，共 {len(parsed_data)} 個 chunks")
            
        except Exception as e:
            logger.error(f"❌ [Parser] YAML 解析失敗: {e}")
        
        return parsed_data
    
    def _parse_xml(self, file_path: str) -> List[Dict[str, Any]]:
        """解析 XML 文件"""
        return self._parse_text(file_path)
    
    def _parse_code(self, file_path: str) -> List[Dict[str, Any]]:
        """解析程式碼文件"""
        parsed_data = []
        file_name = os.path.basename(file_path)
        ext = Path(file_path).suffix.lower()
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            code_text = f"[程式碼文件: {file_name} ({ext})]\n\n{content}"
            
            chunks = self._split_text(code_text)
            for chunk_idx, chunk in enumerate(chunks):
                if len(chunk) > 30:
                    parsed_data.append({
                        "text": chunk,
                        "metadata": {
                            "file_name": file_name,
                            "page_label": "1",
                            "chunk_index": chunk_idx,
                            "content_type": "code",
                            "language": ext.lstrip('.'),
                            "source": "code"
                        }
                    })
            
            logger.info(f"✅ [Parser] 程式碼解析完成，共 {len(parsed_data)} 個 chunks")
            
        except Exception as e:
            logger.error(f"❌ [Parser] 程式碼解析失敗: {e}")
        
        return parsed_data
    
    def _split_text(self, text: str) -> List[str]:
        """將文字切分成適當大小的 chunks"""
        if len(text) <= self.chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + self.chunk_size
            
            if end < len(text):
                for sep in ['\n\n', '\n', '。', '.', '！', '!', '？', '?', '；', ';', '，', ',']:
                    last_sep = text.rfind(sep, start, end)
                    if last_sep > start + self.chunk_size // 2:
                        end = last_sep + 1
                        break
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            start = end - self.chunk_overlap if end < len(text) else end
        
        return chunks


# 全域實例
_parser_instance = None


def get_multimodal_parser() -> MultimodalParser:
    """取得全域解析器實例"""
    global _parser_instance
    if _parser_instance is None:
        _parser_instance = MultimodalParser()
    return _parser_instance
